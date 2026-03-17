import streamlit as st
from google import genai
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
import json
import io
from PIL import Image

# Configuración de la página de Streamlit
st.set_page_config(page_title="Procesador de libros Notariales", page_icon="📝", layout="centered")

st.title("📝 Procesador de Fotos: Libros Notariales")
st.markdown("Esta aplicación extrae el texto de las fotos de los libros notariales (ignorando sellos) y genera un archivo Word formateado.")

# Intentar obtener la API Key de los secretos de Streamlit primero
api_key = st.secrets.get("GEMINI_API_KEY", "")

# Si no está configurada en la nube o se está corriendo en local sin secrets, la pedimos manual
if not api_key:
    api_key = st.text_input("Ingresa tu clave de API de Google Gemini:", type="password", help="Obtén tu API key en https://aistudio.google.com/app/apikey")
else:
    # Mostramos un pequeño mensaje de confirmación sin revelar la clave
    st.success("✅ Clave API detectada automáticamente por el sistema.")

tab1, tab2 = st.tabs(["📁 Subir foto de la galería", "📸 Tomar foto (Recomendado)"])

with tab1:
    uploaded_file_upload = st.file_uploader("Sube una foto del libro notarial (PNG, JPG)", type=["png", "jpg", "jpeg"])

with tab2:
    uploaded_file_camera = st.camera_input("Toma la foto directamente aquí")

# Dar prioridad a la foto tomada con la cámara si existen ambas
uploaded_file = uploaded_file_camera if uploaded_file_camera is not None else uploaded_file_upload

if uploaded_file is not None:
    # Solo mostrar la imagen ampliada si se subió por archivo
    if uploaded_file == uploaded_file_upload:
        st.image(uploaded_file, caption="Vista previa de la imagen", use_column_width=True)

    if not api_key:
        st.warning("Por favor, ingresa tu clave API de Gemini para procesar la imagen.")
    else:
        # 1. Detectar si el usuario seleccionó otra foto distinta para auto-limpiar el caché
        if "last_image" in st.session_state and st.session_state["last_image"] != uploaded_file.name:
            for key in ["docx_data", "docx_name", "json_crudo"]:
                if key in st.session_state:
                    del st.session_state[key]
        st.session_state["last_image"] = uploaded_file.name

        # 2. Si NO tenemos los datos generados, mostrar el botón de procesar
        if "docx_data" not in st.session_state:
            if st.button("Procesar imagen y generar Word", type="primary"):
                with st.spinner("Analizando imagen con Inteligencia Artificial..."):
                    try:
                        client = genai.Client(api_key=api_key)
                        image = Image.open(uploaded_file)
                        
                        prompt = """
                        Eres un asistente experto en transcripción de documentos notariales.
                        Analiza esta imagen y extrae el texto respetando ESTRICTAMENTE estas reglas:
                        
                        1. OMITE Y BORRA el nombre del notario, su título y número de notaría (ej. "LIC. FRANCISCO PORTILLA BONILLA") SOLO cuando aparezca HASTA ARRIBA del documento (antes de "CERTIFICACIÓN NÚMERO...") y HASTA ABAJO (después de "H. CÓRDOBA VER., A 08 DE MARZO..."). 
                        2. IMPORTANTE: SI el nombre del notario ("LICENCIADO FRANCISCO PORTILLA BONILLA...") aparece ADENTRO del cuerpo del texto justificado, AHÍ SÍ DEBES respetarlo y colocarlo tal cual aparece.
                        3. IGNORA cualquier sello que esté encimado en el texto.
                        4. El texto principal a extraer COMIENZA donde dice algo como "CERTIFICACIÓN NÚMERO..." o "LIBRO DE REGISTRO...".
                        5. Para el cuerpo del documento, escribe los párrafos de forma CONTINUA. No pongas saltos de línea a mitad de una oración simplemente porque así se ve en la foto. Un párrafo debe terminar solo cuando haya un punto y aparte.
                        
                        Devuelve ÚNICAMENTE un objeto JSON válido con esta estructura exacta (sin formato markdown de código):
                        {
                            "numero_certificacion": "95993",
                            "encabezado_derecha": [
                                "Línea 1 del encabezado (ej. CERTIFICACIÓN NÚMERO...)",
                                "Línea 2 del encabezado (ej. LIBRO DE REGISTRO...)"
                            ],
                            "cuerpo_justificado": [
                                "Párrafo 1 completo y continuo sin saltos de línea intermedios...",
                                "Párrafo 2 completo y continuo sin saltos de línea intermedios..."
                            ],
                            "cierre_centrado": [
                                "La línea de fecha y lugar final (ej. H. CÓRDOBA VER., A 08 DE MARZO DEL 2023)"
                            ]
                        }
                        Asegúrate de que TODOS los valores en los arrays sean cadenas de texto. No incluyas el nombre del notario en ninguna de las secciones.
                        """
                        
                        modelos_a_probar = ['gemini-2.5-flash', 'gemini-2.0-flash', 'gemini-1.5-flash']
                        response = None
                        ultimo_error = None
                        
                        for nombre_modelo in modelos_a_probar:
                            try:
                                response = client.models.generate_content(
                                    model=nombre_modelo,
                                    contents=[prompt, image]
                                )
                                break
                            except Exception as e:
                                ultimo_error = e
                                continue
                                
                        if response is None:
                            raise Exception(f"Ningún modelo compatible encontrado. Último error: {ultimo_error}")
                        
                        texto_respuesta = response.text.strip()
                        if texto_respuesta.startswith("```json"):
                            texto_respuesta = texto_respuesta[7:]
                        if texto_respuesta.endswith("```"):
                            texto_respuesta = texto_respuesta[:-3]
                            
                        resultado = json.loads(texto_respuesta.strip())
                        
                        num_certificacion = str(resultado.get("numero_certificacion", "Documento"))
                        encabezado_derecha = resultado.get("encabezado_derecha", [])
                        cuerpo_justificado = resultado.get("cuerpo_justificado", [])
                        cierre_centrado = resultado.get("cierre_centrado", [])
                        
                        num_certificacion_limpio = num_certificacion.replace(",", "").replace('"', '').strip()
                        if not num_certificacion_limpio:
                            num_certificacion_limpio = "Documento"
                            
                        # GENERACIÓN DEL DOCUMENTO WORD
                        doc = Document()
                        
                        for linea in encabezado_derecha:
                            p = doc.add_paragraph(linea)
                            p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                            
                        for parrafo in cuerpo_justificado:
                            p = doc.add_paragraph(parrafo)
                            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                            
                        for linea in cierre_centrado:
                            p = doc.add_paragraph(linea)
                            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        
                        docx_io = io.BytesIO()
                        doc.save(docx_io)
                        docx_io.seek(0)
                        
                        nombre_archivo = f"{num_certificacion_limpio}.docx"
                        
                        # Almacenar en caché todo el resultado para persistencia
                        st.session_state["docx_data"] = docx_io.getvalue()
                        st.session_state["docx_name"] = nombre_archivo
                        st.session_state["json_crudo"] = resultado
                        
                        # Forzar la recarga visual limpia
                        st.rerun()
                        
                    except Exception as e:
                        st.error(f"Ocurrió un error procesando la imagen o el formato: {e}")
                        st.info("Asegúrate de que la clave de API sea correcta y de que la foto sea legible.")
        
        # 3. Si YA tenemos el documento generado, mostrar el botón permanente
        if "docx_data" in st.session_state:
            st.success("✨ ¡Análisis completado con éxito! El documento está listo permanentemente para su descarga.")
            
            with st.expander("Ver datos extraídos (JSON Crudo)"):
                st.json(st.session_state["json_crudo"])
                
            # Interfaz mejorada con Columnas para los botones
            col1, col2 = st.columns(2)
            
            with col1:
                st.download_button(
                    label=f"📥 Descargar Documento ({st.session_state['docx_name']})",
                    data=st.session_state["docx_data"],
                    file_name=st.session_state["docx_name"],
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    use_container_width=True
                )
                
                st.info("💡 **Tip para iPhone:** Al tocar Descargar, Safari te preguntará si deseas descargar el archivo. Acéptalo, y luego toca el ícono de **Compartir** (el cuadrito con la flecha hacia arriba) o la lupa azul en la barra del navegador para enviarlo rápido por WhatsApp o Correo.")
            
            with col2:
                # El botón de RESET o Vaciar caché solicitado por el usuario
                if st.button("🔄 Procesar nuevo documento (Reset)", use_container_width=True):
                    for key in ["docx_data", "docx_name", "json_crudo", "last_image"]:
                        if key in st.session_state:
                            del st.session_state[key]
                    st.rerun()
